#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""office-token-booster 执行引擎（方向 B 核心）。

把办公任务直接做成可交付物，并可选自动记回 ledger（执行 + 度量闭环）。

设计红线（保住比赛安全/可审计加分项）：
- 纯标准库，零第三方依赖；不联网、不读密钥、不调 LLM。
- 内容类任务走「模板渲染」：宿主 Agent 提供原始内容，本引擎结构化渲染成 Markdown。
- 数据类任务（CSV）走 stdlib 本地计算，真·本地执行。
- 执行只写用户显式指定的 --output 路径，不扫描技能目录外文件。
- 自动记账复用 ledger_agent.run_long_chain 护栏（baseline 缺省拦截已修），默认 dry-run。
"""
from __future__ import annotations

import argparse
import csv
import html
import importlib
import io
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Tuple

# ---------------------------------------------------------------------------
# 任务类型归一（复用类型字典；缺失则退化为直匹配）
# ---------------------------------------------------------------------------
try:
    from type_registry import load_registry, normalize_type  # type: ignore
    _REGISTRY = load_registry()
except Exception:  # pragma: no cover - 字典缺失时退化
    _REGISTRY = None

    def normalize_type(t: str, registry=None) -> str:  # type: ignore
        return t.strip()


# 执行引擎支持的类型别名（小写匹配）
_EXEC_ALIASES = {
    "周报": "周报生成",
    "周报生成": "周报生成",
    "weekly": "周报生成",
    "weekly report": "周报生成",
    "纪要": "会议纪要",
    "会议纪要": "会议纪要",
    "会议记录": "会议纪要",
    "meeting": "会议纪要",
    "minutes": "会议纪要",
    "数据分析": "数据分析",
    "csv分析": "数据分析",
    "excel分析": "数据分析",
    "数据整理": "数据分析",
    "文档整理": "文档整理",
    "要点提炼": "文档整理",
    "总结": "文档整理",
    "摘要": "文档整理",
    "提炼": "文档整理",
    "ppt": "PPT大纲",
    "ppt大纲": "PPT大纲",
    "幻灯片": "PPT大纲",
    "slides": "PPT大纲",
}


def resolve_exec_type(task_type: str) -> Optional[str]:
    """把用户给的 task_type 归一到执行引擎支持的标准名；不支持则返回 None。"""
    if _REGISTRY is not None:
        try:
            norm = normalize_type(task_type, _REGISTRY)
            if norm in _EXEC_ALIASES.values():
                return norm
        except Exception:
            pass
    key = task_type.strip().lower()
    if key in _EXEC_ALIASES:
        return _EXEC_ALIASES[key]
    # 子串兜底
    for k, v in _EXEC_ALIASES.items():
        if k in key or key in k:
            return v
    return None


# ---------------------------------------------------------------------------
# 通用工具
# ---------------------------------------------------------------------------
def _lines(text: str) -> list[str]:
    return [ln.strip() for ln in text.splitlines() if ln.strip()]


def _today_iso() -> str:
    from datetime import datetime

    return datetime.now().strftime("%Y-%m-%d")


def _detect_date_range(text: str) -> str:
    """从文本里抓 '2026-08-18 ~ 2026-08-24' 这类区间；没有则返回今天。"""
    m = re.search(r"(\d{4}-\d{2}-\d{2})\s*[~–至到-]\s*(\d{4}-\d{2}-\d{2})", text)
    if m:
        return f"{m.group(1)} ~ {m.group(2)}"
    m = re.search(r"(\d{4}-\d{2}-\d{2})", text)
    if m:
        return m.group(1)
    return _today_iso()


# ---------------------------------------------------------------------------
# 模块 1：周报生成
# ---------------------------------------------------------------------------
def render_weekly_report(text: str) -> str:
    lines = _lines(text)
    rng = _detect_date_range(text)
    overview, work, plan, risk = [], [], [], []

    # 行内锚点前缀：一行内混合多个要点时（如「风险：x；下周：y」），按前缀正确路由，
    # 避免整行被「下周」等关键词抢走而丢失「风险」段（E3 增强）。
    # 长标签优先（如「本周概览」「下周计划」），保证去前缀更干净。
    _ANCHORS = [
        (re.compile(r"^\s*(风险|阻塞|卡点|问题|blocker|issue)\s*[:：]?", re.I), risk),
        (re.compile(r"^\s*(下周计划|下週计划|下周|下週|计划|规划|plan|todo|后续|待办)\s*[:：]?", re.I), plan),
        (re.compile(r"^\s*(本周概览|本周|概览|摘要|总览|summary|overview)\s*[:：]?", re.I), overview),
    ]

    def classify(sub: str):
        s = sub.strip()
        if not s:
            return None
        # 1) 行内锚点前缀优先（如「风险：…」「下周计划：…」「概览：…」）
        for pat, bucket in _ANCHORS:
            if pat.match(s):
                content = pat.sub("", s).strip() or s
                return bucket, content
        # 2) 无显式前缀时，按关键词兜底（与旧逻辑一致）
        if re.search(r"下周|下週|计划|规划|plan|todo|后续|待办", s, re.I):
            return plan, s
        if re.search(r"风险|阻塞|卡点|问题|卡住|blocker|issue|待解决", s, re.I):
            return risk, s
        if re.search(r"概览|摘要|本周|总览|summary|overview", s, re.I) and not work:
            return overview, s
        return work, s

    for ln in lines:
        # 一行内多个要点用 ；/ ; 分隔时，拆开分别归类（E3 增强核心）
        for sub in re.split(r"[；;]", ln):
            r = classify(sub)
            if r is None:
                continue
            bucket, content = r
            bucket.append(content)

    out = [f"# 周报（{rng}）", ""]
    if overview:
        out += ["## 本周概览", ""] + [f"- {x}" for x in overview] + [""]
    out += ["## 重点工作", ""]
    out += [f"- {x}" for x in (work or ["（暂无记录）"])] + [""]
    if risk:
        out += ["## 风险与阻塞", ""] + [f"- {x}" for x in risk] + [""]
    out += ["## 下周计划", ""]
    out += [f"- {x}" for x in (plan or ["（待补充）"])] + [""]
    out.append("> 由 office-token-booster 执行引擎生成；节省值请在确认写回账本时补充「笨办法」基准。")
    return "\n".join(out)


# ---------------------------------------------------------------------------
# 模块 2：会议纪要
# ---------------------------------------------------------------------------
_ACTION_RE = re.compile(r"待办|行动|action|todo|跟进|负责人|@|owner", re.I)
_DECISION_RE = re.compile(r"结论|决定|决议|agree|decision|确认", re.I)
_OPEN_RE = re.compile(r"遗留|待定|未决|open|悬而未决", re.I)
_ATTENDEE_RE = re.compile(r"参会|出席|列席|attendee|背景|主题", re.I)


def render_meeting_minutes(text: str) -> str:
    lines = _lines(text)
    rng = _detect_date_range(text)
    attendee, decision, action, open_iss, discuss = [], [], [], [], []
    for ln in lines:
        if _ATTENDEE_RE.search(ln) and not decision and not action:
            attendee.append(ln)
        elif _ACTION_RE.search(ln):
            action.append(ln)
        elif _DECISION_RE.search(ln):
            decision.append(ln)
        elif _OPEN_RE.search(ln):
            open_iss.append(ln)
        else:
            discuss.append(ln)

    out = [f"# 会议纪要（{rng}）", ""]
    if attendee:
        out += ["## 参会与背景", ""] + [f"- {x}" for x in attendee] + [""]
    out += ["## 核心结论", ""]
    out += [f"- {x}" for x in (decision or ["（会议未形成明确结论）"])] + [""]
    out += ["## 待办事项", ""]
    if action:
        for ln in action:
            owner = re.search(r"(负责人[：:]\s*\S+|@\S+)", ln)
            due = re.search(r"(截止[：:]\s*\S+|\d{4}-\d{2}-\d{2})", ln)
            owner_s = f" ｜ 负责人：{owner.group(1)}" if owner else ""
            due_s = f" ｜ 截止：{due.group(1)}" if due else ""
            out.append(f"- [ ] {ln}{owner_s}{due_s}")
    else:
        out.append("- （无明确待办）")
    out += ["", "## 遗留问题", ""]
    out += [f"- {x}" for x in (open_iss or ["（无）"])] + [""]
    if discuss and not attendee:
        out += ["## 其他讨论", ""] + [f"- {x}" for x in discuss] + [""]
    out.append("> 由 office-token-booster 执行引擎生成。")
    return "\n".join(out)


# ---------------------------------------------------------------------------
# 模块 3：数据分析（CSV，纯本地计算）
# ---------------------------------------------------------------------------
def _read_csv_table(text_or_path: str) -> tuple[list[str], list[list[str]]]:
    """返回 (header, rows)。输入是路径则读文件，否则按 CSV 文本解析。"""
    if "\n" not in text_or_path and Path(text_or_path).exists():
        with open(text_or_path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.reader(f)
            rows = [r for r in reader if r]
    else:
        reader = csv.reader(io.StringIO(text_or_path))
        rows = [r for r in reader if r]
    if not rows:
        return [], []
    return rows[0], rows[1:]


def _is_number(s: str) -> bool:
    try:
        float(s)
        return True
    except (ValueError, TypeError):
        return False


def analyze_csv(text_or_path: str) -> str:
    header, rows = _read_csv_table(text_or_path)
    if not header:
        return "# 数据分析\n\n> 未解析到任何数据行。"

    n_cols = len(header)
    n_rows = len(rows)
    col_stats: list[str] = []
    out = [f"# 数据分析报告（{n_rows} 行 × {n_cols} 列）", ""]
    out.append(f"**字段**：{', '.join(header)}")
    out.append("")

    for i, col in enumerate(header):
        vals = [r[i] for r in rows if i < len(r) and r[i] != ""]
        nums = [float(v) for v in vals if _is_number(v)]
        if nums:
            s = sum(nums)
            avg = s / len(nums)
            mx, mn = max(nums), min(nums)
            sorted_n = sorted(nums)
            mid = sorted_n[len(sorted_n) // 2] if sorted_n else 0
            col_stats.append(
                f"| {col} | 数值 | {len(nums)} | {s:,.2f} | {avg:,.2f} | {mn:,.2f} | {mx:,.2f} | {mid:,.2f} |"
            )
        else:
            # 类别列：top 值计数
            from collections import Counter

            top = Counter(vals).most_common(3)
            top_s = "、".join(f"{k}({c})" for k, c in top) or "（空）"
            col_stats.append(f"| {col} | 类别 | {len(vals)} | - | - | - | - | {top_s} |")

    out += ["## 关键指标", "", "| 字段 | 类型 | 非空数 | 求和 | 均值 | 最小 | 最大 | 中位数/Top |", "|---|---|---|---|---|---|---|---|"]
    out += col_stats
    out += ["", "> 由 office-token-booster 执行引擎本地计算（不联网、不读密钥）。"]
    return "\n".join(out)


# ---------------------------------------------------------------------------
# 模块 4：文档整理 / 要点提炼
# ---------------------------------------------------------------------------
def summarize_doc(text: str) -> str:
    # 按 Markdown 标题或空行分段
    blocks = re.split(r"\n\s*\n|(?=^#{1,3}\s)", text.strip(), flags=re.M)
    blocks = [b.strip() for b in blocks if b.strip()]
    outline, points = [], []
    for b in blocks:
        title = b.lstrip("#").strip().splitlines()[0] if b.startswith("#") else b.splitlines()[0][:40]
        outline.append(title)
        # 核心要点：取每段首句（去标题符号）
        body = b.splitlines()
        first_sentence = ""
        for line in body:
            line = line.lstrip("#").strip()
            if len(line) > 4 and not line.startswith(">"):
                first_sentence = re.split(r"[。！？.!?]", line)[0][:80]
                break
        if first_sentence:
            points.append(first_sentence)

    one_liner = points[0] if points else "（文档过短，无可提炼要点）"
    out = ["# 要点提炼", ""]
    out += ["## 文档大纲", ""] + [f"- {x}" for x in outline] + [""]
    out += ["## 核心要点", ""] + [f"- {x}" for x in points] + [""]
    out += ["## 一句话总结", "", one_liner, ""]
    out.append("> 由 office-token-booster 执行引擎生成。")
    return "\n".join(out)


# ---------------------------------------------------------------------------
# 模块 5：PPT 大纲生成
# ---------------------------------------------------------------------------
def render_ppt_outline(text: str) -> str:
    lines = _lines(text)
    title = lines[0] if lines else "未命名主题"
    bullets = lines[1:] if len(lines) > 1 else []
    # 把要点按 ~4 条一组切成内容页
    chunks = [bullets[i : i + 4] for i in range(0, len(bullets), 4)] or [["（补充要点）"]]

    out = ["# 幻灯片大纲", "", f"**主题**：{title}", ""]
    out += ["## Slide 1 · 封面", f"- 标题：{title}", "- 副标题：office-token-booster 执行引擎生成", ""]
    out += ["## Slide 2 · 背景与问题", "- （补充：为什么要讲这个主题）"]
    if bullets:
        out += ["", "## Slide 3 · 核心要点"]
        for b in bullets:
            out.append(f"- {b}")
    out += ["", "## Slide 4 · 数据 / 案例", "- （补充：支撑数据或案例）"]
    out += ["", "## Slide 5 · 总结与行动号召", "- （补充：观众下一步做什么）", ""]
    out.append("> 由 office-token-booster 执行引擎生成；送去排版类 Skill 可直出 PPT。")
    return "\n".join(out)


# ---------------------------------------------------------------------------
# 调度
# ---------------------------------------------------------------------------
_DISPATCH = {
    "周报生成": render_weekly_report,
    "会议纪要": render_meeting_minutes,
    "数据分析": analyze_csv,
    "文档整理": summarize_doc,
    "PPT大纲": render_ppt_outline,
}

# 公开别名：供对话编排层 conversation.py 安全调用（不碰私有 _DISPATCH），
# 也便于测试直接断言支持的类型集合。
EXECUTORS = _DISPATCH


def execute_render(task_type: str, text: str) -> Tuple[bool, str]:
    """按标准类型名渲染交付物。返回 (是否成功, 内容或错误说明)。

    供 conversation 的 execute 意图路由调用；不写盘、不联网、不读密钥，
    纯本地渲染，符合执行层零依赖红线。
    """
    fn = EXECUTORS.get(task_type)
    if not fn:
        return False, f"暂不支持的任务类型：{task_type}（支持：{', '.join(EXECUTORS)}）"
    try:
        return True, fn(text)
    except Exception as e:  # pragma: no cover - 渲染异常兜底
        return False, f"渲染失败：{e}"


def execute(task_type: str, text: str) -> tuple[str, dict]:
    """执行一个任务，返回 (markdown, meta)。task_type 为标准名。"""
    fn = _DISPATCH.get(task_type)
    if fn is None:
        raise ValueError(f"执行引擎不支持的任务类型：{task_type}")
    md = fn(text)
    meta = {"task_type": task_type, "chars": len(md)}
    return md, meta


# ---------------------------------------------------------------------------
# 自动记账闭环（复用 ledger_agent 护栏）
# ---------------------------------------------------------------------------
def propose_ledger(ledger_path: str, task_type: str,
                   cost: Optional[dict] = None,
                   skill_tokens: Optional[int] = None,
                   skill_minutes: Optional[int] = None,
                   baseline_tokens: Optional[int] = None,
                   baseline_minutes: Optional[int] = None,
                   note: Optional[str] = None,
                   apply: bool = False) -> Optional[dict]:
    """执行完后把这笔账记回 ledger。

    cost 为宿主完成事件的用量字典 {"skill_tokens": N, "skill_minutes": M}
    （复用 v0.7 build_completion_event 形态）；显式传入的 skill_tokens/
    skill_minutes 优先于 cost 字典。
    baseline_tokens / baseline_minutes 为用户「笨办法」手搓基准，显式传入可绕过
    P0 护栏直接写回（空账本上不传 baseline 会被护栏拦截，避免负节省污染账本）。
    """
    if cost:
        if skill_tokens is None:
            skill_tokens = cost.get("skill_tokens")
        if skill_minutes is None:
            skill_minutes = cost.get("skill_minutes")
    try:
        from ledger_agent import run_long_chain
    except Exception:
        return None
    return run_long_chain(
        ledger_path, task_type, apply=apply,
        skill_tokens=skill_tokens, skill_minutes=skill_minutes,
        baseline_tokens=baseline_tokens, baseline_minutes=baseline_minutes,
        note=note,
    )


# ---------------------------------------------------------------------------
# 可选 HTML 包装（最小 md→html，仅标题/列表/粗体；用户内容经 escape）
# ---------------------------------------------------------------------------
def _md_to_html(md: str, title: str) -> str:
    out, in_list = [], False
    for raw in md.splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            if in_list:
                out.append("</ul>"); in_list = False
            out.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith("## "):
            if in_list:
                out.append("</ul>"); in_list = False
            out.append(f"<h2>{html.escape(line[3:])}</h2>")
        elif line.startswith("- "):
            if not in_list:
                out.append("<ul>"); in_list = True
            item = html.escape(line[2:])
            item = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", item)
            out.append(f"<li>{item}</li>")
        elif line.startswith("> "):
            if in_list:
                out.append("</ul>"); in_list = False
            out.append(f"<blockquote>{html.escape(line[2:])}</blockquote>")
        elif line.strip() == "":
            if in_list:
                out.append("</ul>"); in_list = False
        else:
            if in_list:
                out.append("</ul>"); in_list = False
            out.append(f"<p>{html.escape(line)}</p>")
    if in_list:
        out.append("</ul>")
    return f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{html.escape(title)}</title></head><body>{''.join(out)}</body></html>"


# ---------------------------------------------------------------------------
# 可选富格式导出（docx / xlsx）—— 可选依赖，缺失则优雅降级，不破坏零依赖默认
# ---------------------------------------------------------------------------
def _try_import(module_name: str):
    """延迟导入可选第三方库；缺失返回 None（调用方据此降级）。"""
    try:
        return importlib.import_module(module_name)
    except Exception:  # pragma: no cover - 依赖缺失时走降级分支
        return None


def _strip_bold(s: str) -> str:
    """去掉 Markdown 粗体标记（docx/xlsx 不渲染 **）。"""
    return s.replace("**", "")


def _is_table_sep(s: str) -> bool:
    """判定一行是否为 Markdown 表格分隔行（|---|---|）。"""
    core = s.replace("|", "").replace("-", "").replace(":", "")
    return core.strip() == ""


def _parse_md_blocks(md: str):
    """把本引擎产出的 Markdown 切成块，供 docx/xlsx 复用。

    块形态：('h1'|'h2'|'h3'|'p'|'quote'|'ul'|'table', payload)。
    - ul 的 payload 为字符串列表；table 的 payload 为二维列表（含表头，分隔行已跳过）。
    """
    blocks = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith("|"):
            tbl: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_str = lines[i].strip()
                if _is_table_sep(row_str):
                    i += 1
                    continue
                tbl.append([c.strip() for c in row_str.strip("|").split("|")])
                i += 1
            if tbl:
                blocks.append(("table", tbl))
            continue
        if s.startswith("# "):
            blocks.append(("h1", s[2:].strip()))
        elif s.startswith("## "):
            blocks.append(("h2", s[3:].strip()))
        elif s.startswith("### "):
            blocks.append(("h3", s[4:].strip()))
        elif s.startswith("- "):
            items: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append(("ul", items))
            continue
        elif s.startswith("> "):
            blocks.append(("quote", s[2:].strip()))
        else:
            blocks.append(("p", s))
        i += 1
    return blocks


def _replace_ext(path: str, ext: str) -> str:
    return str(Path(path).with_suffix(ext))


def _md_to_csv(md: str, csv_path: str) -> None:
    """xlsx 缺失时的零依赖降级：把 Markdown（含表格）写成 CSV。"""
    rows: list[list[str]] = []
    for kind, payload in _parse_md_blocks(md):
        if kind == "table":
            for r in payload:
                rows.append([_strip_bold(c) for c in r])
        elif kind == "ul":
            for it in payload:
                rows.append([_strip_bold(it)])
        else:
            text = _strip_bold(payload if isinstance(payload, str) else " ".join(payload))
            rows.append([text])
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        csv.writer(f).writerows(rows)


def export_docx(md: str, out_path: str, title: str) -> Tuple[str, str]:
    """导出为 .docx（需可选依赖 python-docx）。

    返回 (实际写出路径, 状态)。python-docx 缺失时降级为同路径 .md，
    状态以 'degraded:' 开头；否则状态为 'ok'。
    """
    if _try_import("docx") is None:
        md_path = _replace_ext(out_path, ".md")
        Path(md_path).write_text(md, encoding="utf-8")
        return md_path, "degraded:docx-lib-missing"
    from docx import Document

    doc = Document()
    if title:
        doc.core_properties.title = title
    for kind, payload in _parse_md_blocks(md):
        if kind == "h1":
            doc.add_heading(_strip_bold(payload), level=1)
        elif kind == "h2":
            doc.add_heading(_strip_bold(payload), level=2)
        elif kind == "h3":
            doc.add_heading(_strip_bold(payload), level=3)
        elif kind == "p":
            doc.add_paragraph(_strip_bold(payload))
        elif kind == "quote":
            doc.add_paragraph(_strip_bold(payload), style="Intense Quote")
        elif kind == "ul":
            for it in payload:
                doc.add_paragraph(_strip_bold(it), style="List Bullet")
        elif kind == "table":
            t = doc.add_table(rows=1, cols=len(payload[0]))
            try:
                t.style = "Light Grid Accent 1"
            except Exception:  # pragma: no cover - 样式名随 Word 版本变化
                pass
            hdr = t.rows[0].cells
            for j, c in enumerate(payload[0]):
                hdr[j].text = _strip_bold(c)
            for r in payload[1:]:
                cells = t.add_row().cells
                for j, c in enumerate(r):
                    cells[j].text = _strip_bold(c)
    doc.save(out_path)
    return out_path, "ok"


def export_xlsx(md: str, out_path: str, title: str) -> Tuple[str, str]:
    """导出为 .xlsx（需可选依赖 openpyxl）。

    叙事内容进「内容」sheet，每个 Markdown 表格进独立的「表N」sheet。
    openpyxl 缺失时降级为同路径 .csv（零依赖），状态以 'degraded:' 开头。
    """
    if _try_import("openpyxl") is None:
        csv_path = _replace_ext(out_path, ".csv")
        _md_to_csv(md, csv_path)
        return csv_path, "degraded:xlsx-lib-missing"
    from openpyxl import Workbook

    wb = Workbook()
    content = wb.active
    content.title = "内容"
    if title:
        content.append([title])
    table_idx = 0
    for kind, payload in _parse_md_blocks(md):
        if kind == "table":
            table_idx += 1
            ws = wb.create_sheet(title=f"表{table_idx}")
            for r in payload:
                ws.append([_strip_bold(c) for c in r])
        elif kind in ("h1", "h2", "h3"):
            content.append([_strip_bold(payload)])
        elif kind == "ul":
            for it in payload:
                content.append([_strip_bold(it)])
        else:
            text = _strip_bold(payload if isinstance(payload, str) else " ".join(payload))
            content.append([text])
    wb.save(out_path)
    return out_path, "ok"


FORMATS = ("md", "html", "docx", "xlsx")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def _read_input(src: str) -> str:
    if src == "-":
        return sys.stdin.read()
    p = Path(src)
    if p.exists():
        return p.read_text(encoding="utf-8-sig")
    # 当作纯文本
    return src


def main(argv: Optional[list[str]] = None) -> int:
    ap = argparse.ArgumentParser(description="office-token-booster 执行引擎")
    ap.add_argument("--type", required=False, help="任务类型（周报生成/会议纪要/数据分析/文档整理/PPT大纲）；--init-ledger 时无需此参数")
    ap.add_argument("--input", required=False, help="输入：文件路径、-（stdin）、或直接文本；--init-ledger 时无需此参数")
    ap.add_argument("--output", help="输出路径（缺省打印到 stdout；docx/xlsx 必须指定）")
    ap.add_argument("--format", choices=FORMATS, default="md",
                    help="输出格式：md（默认）/ html / docx / xlsx；docx/xlsx 需可选依赖，缺失自动降级")
    ap.add_argument("--apply-ledger", help="执行后自动记回的 ledger.json 路径")
    ap.add_argument("--skill-tokens", type=int, help="本次执行实测消耗的 Token（来自宿主 event 或自估）")
    ap.add_argument("--skill-minutes", type=int, help="本次执行耗时（分钟）")
    ap.add_argument("--baseline-tokens", type=int, help="本次任务的「笨办法」手搓基准 Token；显式传入可绕过 P0 护栏直接写回")
    ap.add_argument("--baseline-minutes", type=int, help="本次任务的「笨办法」手搓基准耗时（分钟）")
    ap.add_argument("--cost-json", help="宿主完成事件的 cost JSON，如 '{\"skill_tokens\":1800,\"skill_minutes\":5}'，自动合并进记账")
    ap.add_argument("--note", help="记账备注")
    ap.add_argument("--confirm-ledger", action="store_true", help="与 --apply-ledger 同用，真正写回（否则仅预览）")
    ap.add_argument("--init-ledger", help="创建空账本文件（写入 {\"tasks\":[]}），用于首次记账前初始化；指定后无需 --type")
    args = ap.parse_args(argv)

    # D1：一等「创建空账本」命令，替代脆弱的 echo / python -c（Test 2/3 根因）
    if args.init_ledger:
        p = Path(args.init_ledger)
        if p.exists():
            print(f"[提示] 账本已存在，跳过初始化：{args.init_ledger}（如需重建请先删除旧文件）", file=sys.stderr)
            return 0
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(json.dumps({"tasks": []}, ensure_ascii=False), encoding="utf-8")
        print(f"[完成] 已创建空账本：{args.init_ledger}（写入 {{\"tasks\":[]}}）", file=sys.stderr)
        return 0

    if not args.type:
        print("[错误] 缺少 --type（或使用 --init-ledger 创建账本）。", file=sys.stderr)
        return 2

    if not args.input:
        print("[错误] 缺少 --input（或使用 --init-ledger 创建账本）。", file=sys.stderr)
        return 2

    std_type = resolve_exec_type(args.type)
    if std_type is None:
        print(f"[错误] 执行引擎暂不支持「{args.type}」。支持：{', '.join(_DISPATCH)}", file=sys.stderr)
        return 2

    text = _read_input(args.input)
    md, meta = execute(std_type, text)
    title = f"{std_type} · {_today_iso()}"

    # 富格式导出（docx / xlsx）需指定 --output
    if args.format in ("docx", "xlsx"):
        if not args.output:
            print("[错误] --format docx/xlsx 需要配合 --output 指定输出路径", file=sys.stderr)
            return 2
        if args.format == "docx":
            actual, status = export_docx(md, args.output, title)
        else:
            actual, status = export_xlsx(md, args.output, title)
        if status.startswith("degraded"):
            print(f"[完成·降级] 富格式库未安装，已降级导出 → {actual}（{meta['chars']} 字符）", file=sys.stderr)
        else:
            print(f"[完成] 已生成 {std_type} 交付物（{args.format}）→ {actual}（{meta['chars']} 字符）", file=sys.stderr)
    else:
        rendered = _md_to_html(md, title) if args.format == "html" else md
        if args.output:
            out_path = args.output
            if not out_path.endswith((".md", ".html")):
                out_path = out_path + (".html" if args.format == "html" else ".md")
            Path(out_path).write_text(rendered, encoding="utf-8")
            print(f"[完成] 已生成 {std_type} 交付物 → {out_path}（{meta['chars']} 字符）", file=sys.stderr)
        else:
            print(rendered)

    if args.apply_ledger:
        cost = None
        if args.cost_json:
            try:
                cost = json.loads(args.cost_json)
            except Exception as e:
                print(f"[错误] --cost-json 解析失败：{e}", file=sys.stderr)
                return 2
        try:
            res = propose_ledger(
                args.apply_ledger, std_type, cost=cost,
                skill_tokens=args.skill_tokens, skill_minutes=args.skill_minutes,
                baseline_tokens=args.baseline_tokens, baseline_minutes=args.baseline_minutes,
                note=args.note or f"执行引擎：{std_type}", apply=args.confirm_ledger,
            )
        except FileNotFoundError:
            # E4：账本文件不存在时给出友好提示（而非原始栈），交付物已正常生成，仅跳过记账
            print(f"[错误] 记账账本文件不存在：{args.apply_ledger}", file=sys.stderr)
            print("       请先创建空账本后再记账：", file=sys.stderr)
            print(f"       .venv\\Scripts\\python scripts\\executor.py --init-ledger {args.apply_ledger}", file=sys.stderr)
            print("[提示] 本次交付物已正常生成，仅自动记账被跳过。", file=sys.stderr)
            return 0
        if res is None:
            print("[记账] ledger_agent 不可用，跳过自动记账。", file=sys.stderr)
        elif res.get("blocked"):
            block_reason = res.get("block_reason") or res.get("reason") or "（未知原因，请检查 baseline）"
            print(f"[记账] 已拦截：{block_reason}（请补填 baseline 后确认写回）", file=sys.stderr)
        elif args.confirm_ledger:
            print(f"[记账] 已写回账本：{res.get('ledger_path')}", file=sys.stderr)
        else:
            print("[记账] 预览（dry-run）：确认后加 --confirm-ledger 写回。", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
